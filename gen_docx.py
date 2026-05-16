# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor

doc = Document()

with open('g:/team_Software_Engineering_project/实验报告-Web UI接入与前后端分离.md', 'r', encoding='utf-8') as f:
    text = f.read()

sections = re.split(r'\n(?=#{1,3}\s)', text)
for sec in sections:
    sec = sec.strip()
    if not sec:
        continue
    h_match = re.match(r'^(#{1,3})\s+(.+)', sec)
    if h_match:
        level = len(h_match.group(1))
        title = h_match.group(2)
        doc.add_heading(title, level=level)
        body = sec[h_match.end():].strip()
    else:
        body = sec

    for para in body.split('\n\n'):
        para = para.strip()
        if not para or para.startswith('```') or para.startswith('|') or para.startswith('---'):
            continue
        para = re.sub(r'\*\*(.+?)\*\*', r'\1', para)
        para = re.sub(r'`(.+?)`', r'\1', para)
        para = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', para)
        for line in para.split('\n'):
            line = line.strip()
            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(line)

output = 'g:/team_Software_Engineering_project/实验报告-Web_UI_接入与前后端分离.docx'
doc.save(output)
print('Saved:', output)
